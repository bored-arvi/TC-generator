import tkinter as tk
import csv
from docx import Document
import os
from docx.shared import Inches
from PIL import ImageTk, Image
from tkinter.filedialog import askopenfilename, asksaveasfilename
main= tk.Tk()
main.title("project")


def copy_contents(source_file, destination_file):
    try:
        # Open the source and destination files
        source_doc = Document(source_file)
        destination_doc = Document(destination_file)

        # Copy the paragraphs from the source to the destination
        for paragraph in source_doc.paragraphs:
            destination_doc.add_paragraph(paragraph.text)

        # Save the changes in the destination file
        destination_doc.save(destination_file)

        print("Contents copied successfully!")
    except Exception as e:
        print(f"An error occurred: {e}")
def create_word_file(y,z):
    d = {}
    l = [100, 101, "emis_no", 103, "ad_no", 105, 106, 107, 108, 109, 110, 111, 112, 113, 114, 115, 116, 999,
         117, 118, 119, 120000, 121, 122, 123000, 124, 125, 126, 127, 128, 129, 130]
    i, j = 0, 0
    l1=[]
    with open("tc format.csv", "r") as f:
        f1=open("credentials.csv")
        f1r=list(csv.reader(f1))
        print(f1r)
        fr = list(csv.reader(f))
        c = fr[y + 1]
        while i < len(l) and j < len(l):
            if l[i] == 116:
                a = c[j].split(",")
                b = a[:len(a) // 2]
                Str = ""
                for k in range(len(b)):
                    Str = Str + b[k] + ","
                d[l[i]] = Str
                i = i + 1
            elif l[i] == 999:
                a = c[j].split(",")
                b = a[len(a) // 2:]
                Str = ""
                for k in range(len(b)):
                    if k == len(b) - 1:
                        Str = Str + b[k]
                    else:
                        Str = Str + b[k] + ","
                d[l[i]] = Str
                i += 1
                j += 1
            elif str(l[i]) in f1r[0]:
                l1.append([str(l[i]),f1r[0].index(str(l[i]))])
                i+=1
                j+=1
            else:
                d[l[i]] = c[j]
                i = i + 1
                j = j + 1
        for item in l1:
            d[int(item[0])]=f1r[1][item[1]]
        f1.close()
    d=edit_contents(d)
    doc=Document(z)
    for paragraph in doc.paragraphs:
        for placeholder, value in d.items():
            if str(placeholder) in paragraph.text:
                paragraph.text = paragraph.text.replace(str(placeholder), value)
    doc.save(z)
    os.startfile(z)
def check():
    i=Input.get("1.0", "end-1c")
    file=file_name.get()
    output.delete(tk.END)
    f=open("tc format.csv")
    c=list(csv.reader(f))[1:]
    l=[]
    for j in c:
        l.append(j[3])
    f.close()
    print(l)
    for m in range(len(l)):
        if i==l[m]:
            output.insert(tk.END, "please wait")
            print(m)
            d=c[m]
            print(d[-1])
            copy_contents("form format.docx",file)
            create_word_file(m,file)
        else:
            output.insert(tk.END, "Enter correct scholar number")

def edit_contents(d):
    new_window = tk.Toplevel()
    new_window.title("Edit screen")

    def save_changes():
        d[105] = name.get()
        d[106] = mother_name.get()
        d[107] = father_name.get()
        d[119] = present_days.get()
        d[127] = reason.get()
        d[128] = last_date.get()
        d[129] = issue_date.get()
        file = file_name.get()
        new_window.destroy()
        main.destroy()

    tk.Label(new_window, text="Student Name").grid(row=0, column=0, padx=5, pady=5)
    name = tk.Entry(new_window)
    name.grid(row=0, column=1, padx=5, pady=5)
    name.insert(tk.END, d[105])

    tk.Label(new_window, text="Mother's Name").grid(row=1, column=0, padx=5, pady=5)
    mother_name = tk.Entry(new_window)
    mother_name.grid(row=1, column=1, padx=5, pady=5)
    mother_name.insert(tk.END, d[106])

    tk.Label(new_window, text="Father's Name").grid(row=2, column=0, padx=5, pady=5)
    father_name = tk.Entry(new_window)
    father_name.grid(row=2, column=1, padx=5, pady=5)
    father_name.insert(tk.END, d[107])

    tk.Label(new_window, text="No of days present").grid(row=3, column=0, padx=5, pady=5)
    present_days = tk.Entry(new_window)
    present_days.grid(row=3, column=1, padx=5, pady=5)
    present_days.insert(tk.END, d[119])

    tk.Label(new_window, text="Reason for leaving").grid(row=4, column=0, padx=5, pady=5)
    reason = tk.Entry(new_window)
    reason.grid(row=4, column=1, padx=5, pady=5)
    reason.insert(tk.END, d[127])

    tk.Label(new_window, text="Last day in school").grid(row=5, column=0, padx=5, pady=5)
    last_date = tk.Entry(new_window)
    last_date.grid(row=5, column=1, padx=5, pady=5)
    last_date.insert(tk.END, d[128])

    tk.Label(new_window, text="Date of Issue ").grid(row=6, column=0, padx=5, pady=5)
    issue_date= tk.Entry(new_window)
    issue_date.grid(row=6, column=1, padx=5, pady=5)
    issue_date.insert(tk.END, d[129])

    save_button = tk.Button(new_window, text="Save Changes", command=save_changes)
    save_button.grid(row=7, column=0, columnspan=2, padx=5, pady=10)

    new_window.mainloop()
    return d
l1=tk.Label(text="Enter your scholar number ")
l2=tk.Label(text="Transfer certificate")
tk.Label(main, text="File name"). grid(row=2,column=0,padx=5,pady=5)
file_name=tk.Entry(main)
file_name.grid(row=2,column=1,padx=5,pady=5)
file_name.insert(tk.END,"ex.docx")
Input= tk.Text(main,height=1,width=20,bg="light yellow")

button=tk.Button(main,width=10,text="Continue",command=lambda:check(),bd=1)
output=tk.Text(main,height=1,width=20)

l2.grid(row=0,columnspan=2,column=0)
l1.grid(row=1,column=0)
Input.grid(row=1,column=1)
button.grid(row=3,column=1,columnspan=2)
output.grid(row=4,column=1)

tk.mainloop()